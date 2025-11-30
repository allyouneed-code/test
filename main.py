# main.py
import os
import sys
import time
import argparse
from docx import Document

from src.config import app_config
from src.workflows.test_gen_workflow import TestGenerationWorkflow
from src.report import WorkflowReporter, DocxWorkflowReporter

current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# ... (read_file_content 和 read_requirements_from_docx 函数保持原样，此处省略) ...
def read_file_content(filepath: str) -> str:
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"找不到文件: {filepath}")
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def read_requirements_from_docx(filepath: str) -> str:
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"找不到文件: {filepath}")
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip(): full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells if cell.text.strip()]
            if row_cells: full_text.append(" | ".join(row_cells))
    return "\n".join(full_text)

# --- 修改核心入口函数 ---
def run_workflow(code_text: str, requirement_text: str, 
                 req_filename: str = "N/A", 
                 code_filename: str = "N/A", 
                 target_name: str = "未命名对象"):
    """
    运行工作流，支持传入文件来源信息和被测件名称。
    """
    try:
        print(f"启动任务: {target_name}")
        print("正在构建工作流图...")
        workflow_builder = TestGenerationWorkflow(config=app_config)
        app = workflow_builder.build()

        # 初始化状态，存入 GUI 传递的信息
        initial_state = {
            "code": code_text, 
            "requirement": requirement_text,
            
            # --- 新增字段 ---
            "target_name": target_name,       # 被测件名称
            "req_filename": req_filename,     # 需求文件名
            "code_filename": code_filename,   # 代码文件名
            # ---------------
            
            "retry_count": 0,
            "start_time": time.time(), 
            "total_prompt_tokens": 0, "total_completion_tokens": 0, "total_tokens": 0,
            "iteration_history": [], "execution_feedback": "",
            "mutation_test_has_error": False, "mutation_error_details": "",
            "test_failures": 0, "test_errors": 0,
            "evaluation_result": "NOT_STARTED",
            "analysis_report": "", "structured_requirement": "", "generation_prompt": "",
            "test_code": "", "analysis_model": None, "requirement_model": None, "validation_report": ""
        }

        print("\n--- 开始执行分析与生成 ---")
        final_state = app.invoke(initial_state)

        print("\n--- 生成控制台报告 ---")
        reporter = WorkflowReporter(final_state, app_config)
        reporter.generate_report()

        print("\n--- 生成 Word 报告 ---")
        # 文件名加入时间戳防止覆盖
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_docx = f"TestReport_{target_name}_{timestamp}.docx"
        
        docx_reporter = DocxWorkflowReporter(final_state, output_filename=output_docx)
        docx_reporter.generate()
        
        print(f"\n✅ 报告生成成功: {os.path.abspath(output_docx)}")

    except Exception as e:
        print(f"\n❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()

def main():
    # CLI 支持 (可选)
    parser = argparse.ArgumentParser()
    parser.add_argument("req_file")
    parser.add_argument("code_file")
    parser.add_argument("--target", default="CLI_Task", help="被测件名称")
    args = parser.parse_args()

    try:
        req_text = read_requirements_from_docx(args.req_file)
        code_text = read_file_content(args.code_file)
        run_workflow(
            code_text, req_text, 
            req_filename=os.path.basename(args.req_file),
            code_filename=os.path.basename(args.code_file),
            target_name=args.target
        )
    except Exception as e:
        print(e)

if __name__ == "__main__":
    main()
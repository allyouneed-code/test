# src/report/reporter_docx.py
import json
import os
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DocxWorkflowReporter:
    def __init__(self, state: dict, output_filename: str):
        self.state = state
        self.output_filename = output_filename
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(10.5)
        for i in range(1, 4):
            if f'Heading {i}' in self.doc.styles:
                self.doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(46, 116, 181)

    def generate(self):
        print("  -> 正在组装 Word 报告...")
        
        # 1. 标题
        self._add_title()
        
        # 2. 项目输入信息 (新增)
        self._add_project_info()

        # 3. 执行摘要 (紧接其后，不分页)
        self._add_summary()
        
        # 加载数据
        req_data = self._load_json(self.state.get('structured_requirement', '{}'))
        code_data = self._load_json(self.state.get('analysis_report', '{}'))
        
        # 4. 需求模型
        self.doc.add_heading('2. 需求模型分析 (M_req)', level=1)
        if req_data:
            self._add_req_U_unit_under_test(req_data)
            self._add_req_I_interface(req_data)
            self._add_req_B_behavior(req_data)
            self._add_req_C_constraints(req_data)
        else:
            self.doc.add_paragraph("需求模型数据为空。")

        # 5. 代码模型
        self.doc.add_heading('3. 代码模型分析 (M_code)', level=1)
        self.doc.add_heading('3.1 源代码快照', level=2)
        self._add_code_block(self.state.get('code', ''))
        if code_data:
            self._add_code_A_unit_info(code_data)
            self._add_code_S_static_interface(code_data)
            self._add_code_G_control_flow(code_data)

        # 6. 验证与测试
        self._add_validation_section()
        self._add_traceability_matrix()
        self._add_test_code_section()
        
        try:
            self.doc.save(self.output_filename)
            print(f"[Report] Document saved: {self.output_filename}")
        except Exception as e:
            print(f"[Report] Save failed: {e}")

    def _load_json(self, json_str):
        try: return json.loads(json_str)
        except: return None

    # --- 核心修改部分 ---

    def _add_title(self):
        """添加主标题"""
        # 使用被测件名称作为主标题的一部分，或者作为副标题
        target_name = self.state.get('target_name', '未命名项目')
        heading = self.doc.add_heading(f'自动化测试报告: {target_name}', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}").font.color.rgb = RGBColor(128,128,128)
        
        # 关键点：这里不再调用 self.doc.add_page_break()

    def _add_project_info(self):
        """新增：在报告最前方展示输入信息"""
        self.doc.add_heading('0. 测试对象信息', level=1)
        
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # 1. 被测件名称
        self._add_table_row(table, "被测件标识 (Target)", self.state.get('target_name', 'N/A'))
        # 2. 需求文件来源
        self._add_table_row(table, "需求文档来源", self.state.get('req_filename', 'N/A'))
        # 3. 代码文件来源
        self._add_table_row(table, "代码文件来源", self.state.get('code_filename', 'N/A'))
        
        self.doc.add_paragraph("") # 空行分隔

    # --- 以下方法保持原样 (略去部分实现细节以节省篇幅) ---

    def _add_summary(self):
        self.doc.add_heading('1. 执行摘要', level=1)
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        self._add_table_row(table, "最终结果", self.state.get('evaluation_result', 'N/A'))
        self._add_table_row(table, "覆盖率", f"{self.state.get('coverage', 0.0):.2%}")
        self._add_table_row(table, "通过率", f"{self.state.get('pass_rate', 0.0):.2%}")
        self._add_table_row(table, "发现缺陷数 (F-Cases)", str(self.state.get('test_failures', 0)))
        self._add_table_row(table, "执行耗时", f"{self.state.get('total_execution_time', 0):.2f}s")

    def _add_table_row(self, table, col1, col2):
        row = table.add_row()
        row.cells[0].text = str(col1)
        row.cells[1].text = str(col2)

    # ... (其余 _add_req_U_..., _add_code_A_..., 等方法与之前版本一致) ...
    
    # 为了代码完整性，包含必要的辅助方法
    def _set_table_header(self, table, headers):
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            if i < len(hdr_cells):
                hdr_cells[i].text = text
                hdr_cells[i].paragraphs[0].runs[0].bold = True

    def _add_code_block(self, text):
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:fill'), 'F2F2F2')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        p = cell.paragraphs[0]
        p.text = text
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # ================= M_req 部分 =================

    def _add_req_U_unit_under_test(self, data):
        """2.1 被测单元 (U)"""
        self.doc.add_heading('2.1 被测单元 (U - Unit Under Test)', level=2)
        uut = data.get('unit_under_test', {})
        
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        self._add_table_row(table, "标识符 (ID)", uut.get('identifier', 'N/A'))
        self._add_table_row(table, "描述 (Desc)", uut.get('desc', 'N/A'))
        self._add_table_row(table, "需求溯源", str(uut.get('source', 'N/A')))

    def _add_req_I_interface(self, data):
        """2.2 接口规约 (I)"""
        self.doc.add_heading('2.2 接口规约 (I - Interface)', level=2)
        i_spec = data.get('interface_specification', {})
        
        # 输入参数
        self.doc.add_paragraph("输入参数集合:", style='Caption')
        inputs = i_spec.get('input_parameters', [])
        if inputs:
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            self._set_table_header(table, ['参数名', '描述', '约束'])
            for param in inputs:
                constraints = param.get('constraints', [])
                cons_str = "\n".join(constraints) if isinstance(constraints, list) else str(constraints)
                self._add_table_row(table, param.get('name', ''), param.get('desc', ''), cons_str)
        else:
            self.doc.add_paragraph("无输入参数定义。", style='List Bullet')

        # 外部依赖
        deps = i_spec.get('external_dependencies', [])
        if deps:
            self.doc.add_paragraph("外部依赖 (Mock对象):", style='Caption')
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            self._set_table_header(table, ['依赖名称', '交互契约'])
            for dep in deps:
                self._add_table_row(table, dep.get('name', ''), dep.get('contract', ''))

    def _add_req_B_behavior(self, data):
        """2.3 行为模型 (B)"""
        self.doc.add_heading('2.3 行为模型 (B - Behavior)', level=2)
        b_model = data.get('behavioral_model', {})
        
        # 功能场景
        scenarios = b_model.get('functional_scenarios', [])
        if scenarios:
            self.doc.add_paragraph("功能场景 (Functional Scenarios):", style='Heading 3')
            for sc in scenarios:
                self._add_scenario_card(sc, "FUNC")
        
        # 错误场景
        err_scenarios = b_model.get('error_scenarios', [])
        if err_scenarios:
            self.doc.add_paragraph("异常场景 (Error Scenarios):", style='Heading 3')
            for sc in err_scenarios:
                self._add_scenario_card(sc, "ERR")

    def _add_req_C_constraints(self, data):
        """2.4 约束向量 (C)"""
        self.doc.add_heading('2.4 约束向量 (C - Constraints)', level=2)
        
        c_model = data.get('constraints', {})
        param_constraints = c_model.get('param_constraints', []) if isinstance(c_model, dict) else []
        
        if not param_constraints:
            self.doc.add_paragraph("未提取到详细的等价类约束。")
            return

        for pc in param_constraints:
            param_name = pc.get('param_name', 'Unknown')
            self.doc.add_paragraph(f"参数: {param_name}", style='Heading 3')
            
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            self._set_table_header(table, ['类型', '描述', '边界值 (Derived)'])
            
            for valid in pc.get('P_valid', []):
                vals = valid.get('derived_values', [])
                val_str = ", ".join(map(str, vals))
                self._add_table_row(table, "有效 (Valid)", valid.get('description', ''), val_str)
            
            for invalid in pc.get('P_invalid', []):
                vals = invalid.get('derived_values', [])
                val_str = ", ".join(map(str, vals))
                self._add_table_row(table, "无效 (Invalid)", invalid.get('description', ''), val_str)
            
            self.doc.add_paragraph("")

    # ================= M_code 部分 (保持不变) =================

    def _add_code_A_unit_info(self, data):
        """3.2 单元信息 (A)"""
        self.doc.add_heading('3.2 单元信息 (A - Unit Info)', level=2)
        a_model = data.get('A', {})
        
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        self._add_table_row(table, "单元ID", a_model.get('Id', 'N/A'))
        self._add_table_row(table, "类型", a_model.get('Type', 'N/A'))
        self._add_table_row(table, "位置", a_model.get('Location', 'N/A'))

    def _add_code_S_static_interface(self, data):
        """3.3 静态接口 (S)"""
        self.doc.add_heading('3.3 静态接口 (S - Static Interface)', level=2)
        s_model = data.get('S', {})
        
        self.doc.add_paragraph("代码定义的参数:", style='Caption')
        args_in = s_model.get('Arg_in', [])
        if args_in:
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            self._set_table_header(table, ['参数名', '静态类型', '默认值'])
            for arg in args_in:
                self._add_table_row(table, arg.get('name', ''), arg.get('static_type', ''), str(arg.get('default_value', '')))
        else:
            self.doc.add_paragraph("无参数或无法解析参数。", style='List Bullet')

        c_ext = s_model.get('C_ext', [])
        if c_ext:
            self.doc.add_paragraph("检测到的外部调用:", style='Caption')
            for call in c_ext:
                self.doc.add_paragraph(f"调用: {call.get('target_signature', '')}", style='List Bullet')

    def _add_code_G_control_flow(self, data):
        """3.4 控制流与谓词 (G - CFG)"""
        self.doc.add_heading('3.4 控制流与谓词 (G - CFG)', level=2)
        g_model = data.get('G', {})
        
        nodes = g_model.get('Nodes', [])
        edges = g_model.get('Edges', [])
        entry = g_model.get('Entry', 'N/A')
        exit_points = g_model.get('Exit_Points', [])
        
        # 摘要
        p = self.doc.add_paragraph()
        p.add_run(f"基本块数: {len(nodes)} | 跳转边数: {len(edges)} | 入口: {entry}")
        
        # 1. 基本块详情 (Nodes)
        if nodes:
            self.doc.add_paragraph("基本块详情 (Control Flow Nodes):", style='Caption')
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            # 固定列宽以防代码换行太乱
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(1.0) # ID 列
            table.columns[1].width = Inches(5.0) # Code 列

            self._set_table_header(table, ['节点ID', '代码语句 (Statements)'])
            
            for node in nodes:
                node_id = node.get('id', 'N/A')
                stmts = node.get('statements', [])
                # 合并多行代码
                stmt_text = "\n".join(stmts) if isinstance(stmts, list) else str(stmts)
                
                self._add_table_row(table, node_id, stmt_text)
                
                # 设置代码列字体为等宽
                cell = table.rows[-1].cells[1]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
        
        self.doc.add_paragraph("")

        # 2. 跳转边 (Edges)
        if edges:
            self.doc.add_paragraph("关键路径谓词 (Edges & Predicates):", style='Caption')
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            self._set_table_header(table, ['源节点', '目标节点', '跳转条件 (Predicate)'])
            
            for edge in edges:
                pred = edge.get('predicate', 'None')
                # 展示所有跳转关系，有助于理解流程图
                self._add_table_row(table, edge.get('from_node_id'), edge.get('to_node_id'), pred)
        
        # 3. 出口点
        if exit_points:
             self.doc.add_paragraph("")
             self.doc.add_paragraph("出口点 (Exit Points):", style='Caption')
             for ep in exit_points:
                 self.doc.add_paragraph(f"• 节点 {ep.get('node_id')} -> {ep.get('exit_type')}", style='List Bullet')

    # ================= 通用辅助方法 (保持不变) =================

    def _add_validation_section(self):
        """4. 验证结果"""
        self.doc.add_heading('4. 一致性检查结果', level=1)
        val_data = self._load_json(self.state.get('validation_report', '{}'))
        if not val_data:
            self.doc.add_paragraph("无验证数据。")
            return

        # 4.1 缺失
        gaps_req = val_data.get('gaps_req_to_code', [])
        if gaps_req:
            self.doc.add_heading(f'4.1 代码未实现的需求 ({len(gaps_req)})', level=2)
            for gap in gaps_req:
                p = self.doc.add_paragraph()
                run = p.add_run("✖ " + str(gap))
                run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 4.2 冗余
        gaps_code = val_data.get('gaps_code_to_req', [])
        if gaps_code:
            self.doc.add_heading(f'4.2 未定义代码逻辑 ({len(gaps_code)})', level=2)
            for gap in gaps_code:
                p = self.doc.add_paragraph()
                run = p.add_run("? " + str(gap))
                run.font.color.rgb = RGBColor(255, 165, 0)

        # 4.3 匹配
        matches = val_data.get('matches', [])
        if matches:
            self.doc.add_heading(f'4.3 一致项 ({len(matches)})', level=2)
            for m in matches:
                p = self.doc.add_paragraph()
                run = p.add_run("✔ " + str(m))
                run.font.color.rgb = RGBColor(0, 128, 0)

    def _add_test_code_section(self):
        self.doc.add_heading('5. 最终测试代码', level=1)
        self._add_code_block(self.state.get('test_code', ''))

    def _add_summary(self):
        self.doc.add_heading('1. 执行摘要', level=1)
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        self._add_table_row(table, "最终结果", self.state.get('evaluation_result', 'N/A'))
        self._add_table_row(table, "覆盖率", f"{self.state.get('coverage', 0.0):.2%}")
        self._add_table_row(table, "通过率", f"{self.state.get('pass_rate', 0.0):.2%}")
        self._add_table_row(table, "F-Cases", str(self.state.get('test_failures', 0)))
        self._add_table_row(table, "总耗时", f"{self.state.get('total_execution_time', 0):.2f}s")

    def _add_scenario_card(self, sc, type_tag):
        p = self.doc.add_paragraph()
        prefix = "[功能]" if type_tag == "FUNC" else "[异常]"
        color = RGBColor(0, 0, 255) if type_tag == "FUNC" else RGBColor(200, 0, 0)
        
        run = p.add_run(f"{prefix} {sc.get('id', '')}: {sc.get('description', '')}")
        run.bold = True
        run.font.color.rgb = color
        
        detail_text = []
        if 'pre_conditions' in sc: detail_text.append(f"前置: {sc['pre_conditions']}")
        if 'error_conditions' in sc: detail_text.append(f"触发: {sc['error_conditions']}")
        if 'expected_outcome' in sc: detail_text.append(f"预期: {sc['expected_outcome']}")
        
        if detail_text:
            p2 = self.doc.add_paragraph("\n".join(detail_text))
            p2.paragraph_format.left_indent = Inches(0.5)

    def _add_table_row(self, table, *args):
        row_cells = table.add_row().cells
        for i, text in enumerate(args):
            if i < len(row_cells):
                row_cells[i].text = str(text)

    def _set_table_header(self, table, headers):
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            if i < len(hdr_cells):
                hdr_cells[i].text = text
                hdr_cells[i].paragraphs[0].runs[0].bold = True

    def _add_code_block(self, text):
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:fill'), 'F2F2F2')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        p = cell.paragraphs[0]
        p.text = text
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    def _add_traceability_matrix(self):
        """需求追溯矩阵 (RTM)"""
        import ast # 确保导入 ast 模块
        
        self.doc.add_heading('4.3 需求追溯矩阵 (RTM)', level=2)
        self.doc.add_paragraph("下表展示了测试用例与需求场景的覆盖对应关系：")

        # --- 1. 解析测试代码，提取 @pytest.mark.requirement('ID') ---
        mapping = {}
        test_code = self.state.get('test_code', '')
        try:
            tree = ast.parse(test_code)
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef):
                    for decorator in node.decorator_list:
                        # 匹配装饰器
                        if (isinstance(decorator, ast.Call) and 
                            isinstance(decorator.func, ast.Attribute) and 
                            decorator.func.attr == 'requirement' and decorator.args):
                            try:
                                # 兼容不同 Python 版本的 AST 节点类型
                                arg = decorator.args[0]
                                req_id = None
                                if isinstance(arg, ast.Constant): # Python 3.8+
                                    req_id = arg.value
                                elif isinstance(arg, ast.Str):    # Python 3.7
                                    req_id = arg.s
                                
                                if req_id:
                                    if req_id not in mapping: mapping[req_id] = []
                                    mapping[req_id].append(node.name)
                            except: pass
        except Exception as e:
            print(f"解析测试代码失败: {e}")

        # --- 2. 获取需求 ID 和 描述 ---
        req_model = self._load_json(self.state.get('structured_requirement', '{}'))
        req_info = {} # 格式: { "REQ-001": "描述文本..." }
        
        if req_model and 'behavioral_model' in req_model:
            b_model = req_model['behavioral_model']
            # 提取功能场景
            for s in b_model.get('functional_scenarios', []):
                sid = s.get('id')
                if sid: req_info[sid] = s.get('description', 'N/A')
            # 提取异常场景
            for s in b_model.get('error_scenarios', []):
                sid = s.get('id')
                if sid: req_info[sid] = s.get('description', 'N/A')

        # --- 3. 绘制 Word 表格 (4列) ---
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        table.autofit = False 
        table.columns[0].width = Inches(1.2) # ID
        table.columns[1].width = Inches(2.8) # 描述 (最宽)
        table.columns[2].width = Inches(0.8) # 状态
        table.columns[3].width = Inches(1.5) # 测试用例

        self._set_table_header(table, ['需求ID', '需求场景描述', '状态', '关联用例'])
        
        # 排序并填充
        for rid in sorted(req_info.keys()):
            desc = req_info[rid]
            test_cases = mapping.get(rid, [])
            
            status = "已覆盖" if test_cases else "缺失"
            cases_str = "\n".join(test_cases) if test_cases else "---" 
            
            self._add_table_row(table, rid, desc, status, cases_str)
            
        # --- 4. 检查幻觉 ID ---
        extra_ids = set(mapping.keys()) - set(req_info.keys())
        if extra_ids:
            self.doc.add_paragraph("") # 空行
            self.doc.add_paragraph("[警告] 代码中发现了未定义的需求ID:", style='Caption')
            table_extra = self.doc.add_table(rows=1, cols=3)
            table_extra.style = 'Table Grid'
            self._set_table_header(table_extra, ['未知ID', '描述', '关联用例'])
            for rid in extra_ids:
                self._add_table_row(table_extra, rid, "未知/未在模型中定义", "\n".join(mapping[rid]))